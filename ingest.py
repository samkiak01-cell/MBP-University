"""
MBP University — Document Ingestion & Embedding Pipeline

Parses .docx SOP files and .xlsx FAQ files, chunks them, embeds them,
and returns an in-memory FAISS index.  Nothing is written to disk —
the index lives in RAM and is cached by Streamlit's @st.cache_resource.
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional

import faiss
import numpy as np
import openpyxl
from docx import Document
from sentence_transformers import SentenceTransformer

from config import (
    DOCUMENTS_DIR,
    EMBEDDING_MODEL,
    EMBEDDING_DIMENSION,
    CHUNK_MAX_TOKENS,
    CHUNK_OVERLAP_TOKENS,
)

logger = logging.getLogger(__name__)


# ────────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────────

def _approx_token_count(text: str) -> int:
    """Rough token estimate (≈ words × 1.3)."""
    return int(len(text.split()) * 1.3)


def _is_section_heading(paragraph) -> bool:
    """
    Detect SOP section headings: bold text that looks like a title.
    SOPs use bold List-Paragraph items instead of Word heading styles.
    """
    text = paragraph.text.strip()
    if not text or len(text) > 120:
        return False

    runs_with_text = [r for r in paragraph.runs if r.text.strip()]
    if not runs_with_text:
        return False

    all_bold = all(r.bold for r in runs_with_text)
    if not all_bold:
        return False

    # ── Skip generic one-word / role labels ─────────────────────────────
    clean = text.rstrip(":- ")
    _GENERIC_LABELS = {
        "steps", "purpose", "owner", "timing", "when applicable",
        "contact history", "cs attempts", "collection attempts",
        "requested assistance", "common triggers", "summary of call",
        "accounting impact", "interest application",
    }
    if clean.lower() in _GENERIC_LABELS:
        return False
    if clean.lower().startswith("owner:"):
        return False
    if clean.lower().startswith("timing:"):
        return False

    # Skip lines that look like dialogue / email templates
    if any(p in text.lower() for p in ["hi [", "this is [", "thank you", "calling from"]):
        return False

    # ── Positive signals ────────────────────────────────────────────────
    # Mostly-uppercase titles like "PURPOSE", "SCOPE"
    upper_ratio = sum(1 for c in text if c.isupper()) / max(len(text.replace(" ", "")), 1)
    if upper_ratio > 0.6:
        return True

    # Short bold multi-word text → likely a sub-heading
    if len(text) < 80 and len(clean.split()) >= 2:
        return True

    return False


# ────────────────────────────────────────────────────────────────────────
# DOCX Parsing
# ────────────────────────────────────────────────────────────────────────

def _extract_table_text(table) -> str:
    """Extract all text from a docx table as a readable block."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def parse_docx(filepath: str) -> list[dict]:
    """
    Parse a .docx SOP into sections using bold-heading detection.

    Returns list of dicts:
        { "text", "section", "source", "type": "sop" }
    """
    doc = Document(filepath)
    filename = os.path.basename(filepath)
    sections: list[dict] = []
    current_section = "Introduction"
    current_text_parts: list[str] = []

    def _flush():
        text = "\n".join(current_text_parts).strip()
        if text:
            sections.append({
                "text": text,
                "section": current_section,
                "source": filename,
                "type": "sop",
            })

    # Walk body XML to interleave paragraphs and tables in document order
    from docx.oxml.ns import qn

    body = doc.element.body
    para_idx = 0
    table_idx = 0
    paragraphs = doc.paragraphs
    tables = doc.tables

    for child in body:
        if child.tag == qn("w:p"):
            if para_idx < len(paragraphs):
                p = paragraphs[para_idx]
                para_idx += 1
                if _is_section_heading(p):
                    _flush()
                    current_section = p.text.strip()
                    current_text_parts = [f"## {current_section}"]
                elif p.text.strip():
                    current_text_parts.append(p.text.strip())
        elif child.tag == qn("w:tbl"):
            if table_idx < len(tables):
                tbl_text = _extract_table_text(tables[table_idx])
                table_idx += 1
                if tbl_text.strip():
                    current_text_parts.append(tbl_text)

    _flush()
    return sections


# ────────────────────────────────────────────────────────────────────────
# XLSX (FAQ) Parsing
# ────────────────────────────────────────────────────────────────────────

def parse_xlsx(filepath: str) -> list[dict]:
    """
    Parse the FAQ Excel file.  Each row → one chunk.

    Returns list of dicts with keys:
        text, source, type, topic, location, question, resource_link, section
    """
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    ws = wb.active

    # Read header row
    headers = [
        str(c.value).strip().lower().rstrip() if c.value else ""
        for c in next(ws.iter_rows(min_row=1, max_row=1))
    ]

    col_map: dict[str, int] = {}
    for i, h in enumerate(headers):
        if "topic" in h:
            col_map["topic"] = i
        elif "location" in h:
            col_map["location"] = i
        elif "question" in h:
            col_map["question"] = i
        elif "answer" in h:
            col_map["answer"] = i
        elif "resource" in h or "link" in h:
            col_map["resource_link"] = i

    chunks: list[dict] = []
    filename = os.path.basename(filepath)

    for row in ws.iter_rows(min_row=2, values_only=True):
        question = str(row[col_map.get("question", 2)] or "").strip()
        answer = str(row[col_map.get("answer", 3)] or "").strip()
        if not question:
            continue

        topic = str(row[col_map.get("topic", 0)] or "").strip()
        location = str(row[col_map.get("location", 1)] or "").strip()
        resource = str(row[col_map.get("resource_link", 4)] or "").strip()

        parts = []
        if topic:
            parts.append(f"Topic: {topic}")
        if location:
            parts.append(f"Location: {location}")
        parts.append(f"Q: {question}")
        parts.append(f"A: {answer}")
        text = "\n".join(parts)

        chunks.append({
            "text": text,
            "source": filename,
            "type": "faq",
            "topic": topic,
            "location": location,
            "question": question,
            "resource_link": resource,
            "section": f"FAQ — {topic}",
        })

    wb.close()
    return chunks


# ────────────────────────────────────────────────────────────────────────
# Chunking (split large SOP sections)
# ────────────────────────────────────────────────────────────────────────

def _split_large_chunk(
    chunk: dict,
    max_tokens: int = CHUNK_MAX_TOKENS,
    overlap_tokens: int = CHUNK_OVERLAP_TOKENS,
) -> list[dict]:
    """Split an SOP section that exceeds max_tokens into overlapping sub-chunks."""
    text = chunk["text"]
    if _approx_token_count(text) <= max_tokens:
        return [chunk]

    sentences = re.split(r"(?<=[.!?])\s+", text)
    sub_chunks: list[dict] = []
    current_sentences: list[str] = []
    current_tokens = 0

    section_prefix = f"## {chunk['section']}\n" if chunk.get("section") else ""

    for sent in sentences:
        sent_tokens = _approx_token_count(sent)
        if current_tokens + sent_tokens > max_tokens and current_sentences:
            sub_text = section_prefix + " ".join(current_sentences)
            sub_chunks.append({**chunk, "text": sub_text})

            # Keep overlap
            overlap_sents: list[str] = []
            overlap_count = 0
            for s in reversed(current_sentences):
                t = _approx_token_count(s)
                if overlap_count + t > overlap_tokens:
                    break
                overlap_sents.insert(0, s)
                overlap_count += t
            current_sentences = overlap_sents
            current_tokens = overlap_count

        current_sentences.append(sent)
        current_tokens += sent_tokens

    if current_sentences:
        sub_text = section_prefix + " ".join(current_sentences)
        sub_chunks.append({**chunk, "text": sub_text})

    return sub_chunks


# ────────────────────────────────────────────────────────────────────────
# Discovery & Ingestion
# ────────────────────────────────────────────────────────────────────────

def discover_documents(docs_dir: str = DOCUMENTS_DIR) -> tuple[list[str], list[str]]:
    """Return (docx_paths, xlsx_paths) found in the documents directory."""
    docs_path = Path(docs_dir)
    if not docs_path.exists():
        logger.warning(f"Documents directory does not exist: {docs_dir}")
        return [], []

    docx_files = sorted(str(f) for f in docs_path.glob("*.docx") if not f.name.startswith("~$"))
    xlsx_files = sorted(str(f) for f in docs_path.glob("*.xlsx") if not f.name.startswith("~$"))

    logger.info(f"Discovered {len(docx_files)} .docx and {len(xlsx_files)} .xlsx in {docs_dir}")
    return docx_files, xlsx_files


def ingest_all(docs_dir: str = DOCUMENTS_DIR) -> tuple[list[dict], list[str]]:
    """Parse and chunk every document.  Returns (chunks, filenames)."""
    docx_files, xlsx_files = discover_documents(docs_dir)
    all_chunks: list[dict] = []
    all_filenames: list[str] = []

    for fpath in docx_files:
        logger.info(f"Parsing DOCX: {fpath}")
        try:
            sections = parse_docx(fpath)
            for sec in sections:
                all_chunks.extend(_split_large_chunk(sec))
            all_filenames.append(os.path.basename(fpath))
        except Exception as e:
            logger.error(f"Failed to parse {fpath}: {e}", exc_info=True)

    for fpath in xlsx_files:
        logger.info(f"Parsing XLSX: {fpath}")
        try:
            faq_chunks = parse_xlsx(fpath)
            all_chunks.extend(faq_chunks)
            all_filenames.append(os.path.basename(fpath))
        except Exception as e:
            logger.error(f"Failed to parse {fpath}: {e}", exc_info=True)

    logger.info(f"Ingestion complete: {len(all_chunks)} chunks from {len(all_filenames)} files")
    return all_chunks, all_filenames


def build_vector_store(
    chunks: list[dict],
    model: SentenceTransformer,
) -> faiss.IndexFlatIP:
    """Embed all chunks and return an in-memory FAISS index."""
    texts = [c["text"] for c in chunks]
    logger.info(f"Embedding {len(texts)} chunks …")
    embeddings = model.encode(texts, show_progress_bar=False, normalize_embeddings=True)
    embeddings = np.array(embeddings, dtype="float32")

    index = faiss.IndexFlatIP(EMBEDDING_DIMENSION)
    index.add(embeddings)
    logger.info(f"FAISS index built with {index.ntotal} vectors")
    return index
